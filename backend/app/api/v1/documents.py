from io import BytesIO
import os
import shutil
import uuid
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from app.core.config import get_settings
from app.core.db import get_db
from app.models.db_models import DocumentDB
from app.models.document import Document
from app.models.response import APIResponse
from app.services.platform_service import get_or_create_default_user, log_activity

router = APIRouter()

ALLOWED_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}


class DocumentPreviewResponse(BaseModel):
    document: Document
    preview_url: str


@router.post("/upload", response_model=APIResponse[Document])
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> APIResponse[Document]:
    settings = get_settings()
    mime_type = file.content_type or ""
    extension = os.path.splitext(file.filename or "")[1].lower()
    if mime_type not in ALLOWED_TYPES and extension not in ALLOWED_TYPES.values():
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, and TXT files are supported")

    contents = await file.read()
    user = get_or_create_default_user(db)
    stored_name = f"{uuid.uuid4()}-{file.filename}"
    storage_path = os.path.join(settings.DOCUMENTS_DIR, stored_name)
    with open(storage_path, "wb") as buffer:
        buffer.write(contents)

    preview_text = _extract_text(contents, mime_type=mime_type, filename=file.filename or "")
    document = DocumentDB(
        user_id=user.id,
        filename=file.filename or stored_name,
        storage_path=storage_path,
        size=len(contents),
        mime_type=mime_type or _guess_mime(extension),
        extracted_text=preview_text,
        metadata_json={"original_filename": file.filename},
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    log_activity(
        db,
        user.id,
        event_type="document_upload",
        entity_type="document",
        entity_id=str(document.id),
        details={"filename": document.filename, "size": document.size},
    )
    return APIResponse(success=True, data=_to_document(document), message="Document uploaded")


@router.get("/", response_model=APIResponse[list[Document]])
async def list_documents(db: Session = Depends(get_db)) -> APIResponse[list[Document]]:
    user = get_or_create_default_user(db)
    documents = (
        db.query(DocumentDB)
        .filter(DocumentDB.user_id == user.id)
        .order_by(DocumentDB.uploaded_at.desc())
        .all()
    )
    return APIResponse(success=True, data=[_to_document(item) for item in documents])


@router.get("/{document_id}", response_model=APIResponse[DocumentPreviewResponse])
async def get_document(document_id: uuid.UUID, db: Session = Depends(get_db)) -> APIResponse[DocumentPreviewResponse]:
    document = db.query(DocumentDB).filter(DocumentDB.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    preview = DocumentPreviewResponse(
        document=_to_document(document),
        preview_url=f"/api/v1/documents/{document.id}/preview",
    )
    return APIResponse(success=True, data=preview)


@router.get("/{document_id}/preview")
async def preview_document(document_id: uuid.UUID, db: Session = Depends(get_db)):
    document = db.query(DocumentDB).filter(DocumentDB.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return FileResponse(document.storage_path, media_type=document.mime_type, filename=document.filename)


@router.patch("/{document_id}", response_model=APIResponse[Document])
async def replace_document(
    document_id: uuid.UUID,
    file: UploadFile | None = File(default=None),
    filename: str | None = Form(default=None),
    db: Session = Depends(get_db),
) -> APIResponse[Document]:
    document = db.query(DocumentDB).filter(DocumentDB.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    if filename:
        document.filename = filename

    if file is not None:
        mime_type = file.content_type or ""
        extension = os.path.splitext(file.filename or "")[1].lower()
        if mime_type not in ALLOWED_TYPES and extension not in ALLOWED_TYPES.values():
            raise HTTPException(status_code=400, detail="Only PDF, DOCX, and TXT files are supported")
        contents = await file.read()
        with open(document.storage_path, "wb") as buffer:
            buffer.write(contents)
        document.filename = filename or file.filename or document.filename
        document.size = len(contents)
        document.mime_type = mime_type or _guess_mime(extension)
        document.extracted_text = _extract_text(contents, mime_type=document.mime_type, filename=document.filename)
        document.metadata_json = {**(document.metadata_json or {}), "replaced": True}

    db.commit()
    db.refresh(document)
    user = get_or_create_default_user(db)
    log_activity(
        db,
        user.id,
        event_type="document_replace",
        entity_type="document",
        entity_id=str(document.id),
        details={"filename": document.filename},
    )
    return APIResponse(success=True, data=_to_document(document), message="Document updated")


@router.delete("/{document_id}", response_model=APIResponse[bool])
async def delete_document(document_id: uuid.UUID, db: Session = Depends(get_db)) -> APIResponse[bool]:
    document = db.query(DocumentDB).filter(DocumentDB.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    if os.path.exists(document.storage_path):
        os.remove(document.storage_path)
    user_id = document.user_id
    db.delete(document)
    db.commit()
    log_activity(
        db,
        user_id,
        event_type="document_delete",
        entity_type="document",
        entity_id=str(document_id),
    )
    return APIResponse(success=True, data=True, message="Document deleted")


def _to_document(document: DocumentDB) -> Document:
    return Document(
        id=document.id,
        user_id=str(document.user_id),
        filename=document.filename,
        size=document.size,
        mime_type=document.mime_type,
        uploaded_at=document.uploaded_at,
        preview_text=(document.extracted_text or "")[:1200],
        metadata=document.metadata_json or {},
        created_at=document.created_at,
        updated_at=document.updated_at,
    )


def _extract_text(contents: bytes, mime_type: str, filename: str) -> str:
    extension = os.path.splitext(filename)[1].lower()
    if mime_type == "text/plain" or extension == ".txt":
        return contents.decode("utf-8", errors="ignore")

    if mime_type == "application/pdf" or extension == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(contents))
            return "\n".join((page.extract_text() or "") for page in reader.pages)[:15000]
        except Exception:
            return ""

    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or extension == ".docx":
        try:
            from docx import Document as DocxDocument

            document = DocxDocument(BytesIO(contents))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)[:15000]
        except Exception:
            return ""

    return ""


def _guess_mime(extension: str) -> str:
    for mime_type, allowed_extension in ALLOWED_TYPES.items():
        if allowed_extension == extension:
            return mime_type
    return "application/octet-stream"
